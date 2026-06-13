from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path(r"D:\laragon\www\ql_su_kien")
SOURCE_DOCX = Path(
    r"C:\Users\phuqu\OneDrive\Máy tính\64131942_DuongPhuQuang_DoAn - Sao chép - viet tat trung lap highlight.docx"
)
WORKSPACE_OUT = WORKSPACE / "64131942_DuongPhuQuang_DoAn_powerdesigner_luocbot_trichnguon_highlight.docx"
DESKTOP_OUT = Path(
    r"C:\Users\phuqu\OneDrive\Máy tính\64131942_DuongPhuQuang_DoAn - Sao chép - PowerDesigner luoc bot trich nguon highlight.docx"
)
FIG_DIR = WORKSPACE / "docx_powerdesigner_style_figures"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\tahoma.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


FONT_16 = font(16)
FONT_18 = font(18)
FONT_20 = font(20)
FONT_22_B = font(22, bold=True)
FONT_26_B = font(26, bold=True)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, text_font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        candidate = word if not line else f"{line} {word}"
        bbox = draw.textbbox((0, 0), candidate, font=text_font)
        if bbox[2] - bbox[0] <= max_width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    text_font,
    fill: str = "#111827",
    line_gap: int = 4,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, text_font, max(20, x2 - x1 - 20))
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=text_font)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + max(0, len(lines) - 1) * line_gap
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=text_font, fill=fill)
        y += h + line_gap


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    label: str,
    fill: str,
    outline: str,
    text_fill: str = "#111827",
    radius: int = 10,
    width: int = 2,
    text_font=FONT_18,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    centered_text(draw, box, label, text_font, text_fill)


def process_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    number: str,
    label: str,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=14, fill="#DBEAFE", outline="#2563EB", width=3)
    draw.ellipse((x1 + 10, y1 + 10, x1 + 48, y1 + 48), fill="#2563EB", outline="#1D4ED8")
    centered_text(draw, (x1 + 10, y1 + 10, x1 + 48, y1 + 48), number, FONT_16, "white")
    centered_text(draw, (x1 + 54, y1 + 8, x2 - 8, y2 - 8), label, FONT_18, "#0F172A")


def datastore(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    label: str,
    code: str,
) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#FEF3C7", outline="#B45309", width=2)
    draw.line((x1 + 12, y1, x1 + 12, y2), fill="#B45309", width=2)
    draw.line((x2 - 12, y1, x2 - 12, y2), fill="#B45309", width=2)
    draw.text((x1 + 20, y1 + 7), code, font=FONT_16, fill="#92400E")
    centered_text(draw, (x1 + 18, y1 + 24, x2 - 18, y2 - 6), label, FONT_16, "#111827")


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str = "",
    color: str = "#334155",
    width: int = 2,
    label_offset: tuple[int, int] = (0, -18),
) -> None:
    draw.line((start, end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 11
    p1 = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    p2 = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, p1, p2], fill=color)
    if label:
        mid = ((start[0] + end[0]) / 2 + label_offset[0], (start[1] + end[1]) / 2 + label_offset[1])
        bbox = draw.textbbox((0, 0), label, font=FONT_16)
        pad = 4
        draw.rounded_rectangle(
            (mid[0] - pad, mid[1] - pad, mid[0] + bbox[2] - bbox[0] + pad, mid[1] + bbox[3] - bbox[1] + pad),
            radius=4,
            fill="#FFFFFF",
            outline="#CBD5E1",
        )
        draw.text(mid, label, font=FONT_16, fill="#111827")


def new_canvas(title: str, w: int = 1800, h: int = 1120) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, w, 80), fill="#F8FAFC")
    draw.text((40, 23), title, font=FONT_26_B, fill="#0F172A")
    draw.text((w - 355, 30), "DFD - dạng PowerDesigner", font=FONT_18, fill="#475569")
    return img, draw


def make_context_diagram(path: Path) -> None:
    img, draw = new_canvas("DFD mức ngữ cảnh - Hệ thống quản lý sự kiện", 1700, 900)
    process_box(draw, (610, 330, 1090, 540), "0", "Hệ thống quản lý sự kiện")
    rounded_box(draw, (80, 190, 370, 310), "Sinh viên", "#F1F5F9", "#64748B", text_font=FONT_20)
    rounded_box(draw, (80, 580, 370, 700), "Quản trị viên", "#F1F5F9", "#64748B", text_font=FONT_20)
    rounded_box(draw, (1330, 190, 1620, 310), "Cán bộ / Khoa", "#F1F5F9", "#64748B", text_font=FONT_20)
    rounded_box(draw, (1330, 580, 1620, 700), "Dịch vụ Email / QR", "#F1F5F9", "#64748B", text_font=FONT_20)
    arrow(draw, (370, 250), (610, 390), "Đăng nhập, đăng ký, điểm danh", label_offset=(-130, -45))
    arrow(draw, (370, 615), (610, 455), "Quản lý danh mục, sự kiện", label_offset=(-140, -48))
    arrow(draw, (610, 520), (370, 675), "Thông báo, kết quả điểm", label_offset=(-130, 28))
    arrow(draw, (1090, 390), (1330, 250), "Báo cáo, thống kê", label_offset=(0, -45))
    arrow(draw, (1330, 650), (1090, 495), "SMTP, mã QR", label_offset=(10, 8))
    arrow(draw, (1090, 455), (1330, 640), "Email, mã điểm danh", label_offset=(10, -8))
    img.save(path)


def make_level1_diagram(path: Path) -> None:
    img, draw = new_canvas("DFD mức 1 - Các tiến trình chính", 2200, 1550)
    externals = {
        "student": (60, 175, 330, 275, "Sinh viên"),
        "admin": (60, 905, 330, 1005, "Quản trị viên"),
        "email": (1840, 505, 2135, 605, "Dịch vụ Email / QR"),
    }
    for box in externals.values():
        rounded_box(draw, box[:4], box[4], "#F1F5F9", "#64748B", text_font=FONT_18)

    processes = [
        ((455, 130, 785, 240), "1.0", "Xác thực người dùng"),
        ((935, 130, 1265, 240), "2.0", "Quản lý sự kiện"),
        ((1415, 130, 1745, 240), "3.0", "Đăng ký / hủy đăng ký"),
        ((455, 520, 785, 630), "4.0", "Điểm danh QR"),
        ((935, 520, 1265, 630), "5.0", "Điểm rèn luyện"),
        ((1415, 520, 1745, 630), "6.0", "Thông báo / email"),
        ((455, 910, 785, 1020), "7.0", "Báo cáo"),
        ((935, 910, 1265, 1020), "8.0", "Bầu cử"),
        ((1415, 910, 1745, 1020), "9.0", "Media / chatbot"),
    ]
    for box, number, label in processes:
        process_box(draw, box, number, label)

    stores = [
        ((455, 300, 785, 385), "D1", "nguoi_dung"),
        ((850, 300, 1085, 385), "D2", "su_kien"),
        ((1115, 300, 1350, 385), "D11", "loai_su_kien"),
        ((1415, 300, 1745, 385), "D3", "dang_ky"),
        ((455, 690, 785, 775), "D4", "chi_tiet_diem_danh"),
        ((935, 690, 1265, 775), "D5", "lich_su_diem"),
        ((1365, 690, 1595, 775), "D6", "thong_bao"),
        ((1625, 690, 1855, 775), "D9", "smtp_settings"),
        ((455, 1080, 785, 1165), "D10", "activity_logs"),
        ((935, 1080, 1265, 1165), "D7", "bau_cu"),
        ((1415, 1080, 1745, 1165), "D8", "thu_vien_da_phuong_tien"),
    ]
    for box, code, label in stores:
        datastore(draw, box, label, code)

    arrow(draw, (330, 225), (455, 185), "đăng nhập", label_offset=(-36, -34))
    arrow(draw, (785, 185), (935, 185), "phiên", label_offset=(-8, -34))
    arrow(draw, (1265, 185), (1415, 185), "dữ liệu sự kiện", label_offset=(-40, -34))
    arrow(draw, (1580, 240), (1580, 520), "đăng ký", label_offset=(18, 30))
    arrow(draw, (1415, 575), (1265, 575), "kết quả", label_offset=(-34, -34))
    arrow(draw, (620, 630), (620, 910), "dữ liệu điểm danh", label_offset=(18, 75))
    arrow(draw, (1100, 630), (1100, 910), "điểm", label_offset=(18, 75))
    arrow(draw, (1580, 630), (1580, 910), "thông báo", label_offset=(18, 75))
    arrow(draw, (330, 955), (455, 965), "quản trị", label_offset=(-24, -36))
    arrow(draw, (1745, 575), (1840, 555), "email/QR", label_offset=(-18, -34))

    arrow(draw, (620, 240), (620, 300), "", label_offset=(0, 0))
    arrow(draw, (1060, 240), (970, 300), "", label_offset=(0, 0))
    arrow(draw, (1140, 240), (1232, 300), "", label_offset=(0, 0))
    arrow(draw, (1580, 240), (1580, 300), "", label_offset=(0, 0))
    arrow(draw, (620, 630), (620, 690), "", label_offset=(0, 0))
    arrow(draw, (1100, 630), (1100, 690), "", label_offset=(0, 0))
    arrow(draw, (1530, 630), (1480, 690), "", label_offset=(0, 0))
    arrow(draw, (1640, 630), (1740, 690), "", label_offset=(0, 0))
    arrow(draw, (620, 1020), (620, 1080), "", label_offset=(0, 0))
    arrow(draw, (1100, 1020), (1100, 1080), "", label_offset=(0, 0))
    arrow(draw, (1580, 1020), (1580, 1080), "", label_offset=(0, 0))

    draw.rounded_rectangle((420, 1265, 1780, 1355), radius=8, fill="#F8FAFC", outline="#CBD5E1", width=2)
    centered_text(
        draw,
        (440, 1276, 1760, 1344),
        "Các tiến trình báo cáo, bầu cử và media có thể đọc dữ liệu từ sự kiện, đăng ký, điểm danh và điểm rèn luyện theo nhu cầu nghiệp vụ.",
        FONT_18,
        "#334155",
    )
    img.save(path)


def make_registration_diagram(path: Path) -> None:
    img, draw = new_canvas("DFD mức 2 - Đăng ký / hủy đăng ký sự kiện", 1800, 1120)
    rounded_box(draw, (70, 185, 330, 285), "Sinh viên", "#F1F5F9", "#64748B", text_font=FONT_20)
    steps = [
        ((510, 150, 840, 260), "3.1", "Tra cứu sự kiện"),
        ((510, 385, 840, 495), "3.2", "Kiểm tra điều kiện đăng ký"),
        ((510, 620, 840, 730), "3.3", "Ghi nhận đăng ký / hủy"),
        ((510, 855, 840, 965), "3.4", "Gửi thông báo kết quả"),
    ]
    for box, number, label in steps:
        process_box(draw, box, number, label)
    stores = [
        ((1140, 150, 1450, 240), "D2", "su_kien"),
        ((1140, 290, 1450, 380), "D3", "dang_ky"),
        ((1140, 430, 1450, 520), "D1", "nguoi_dung"),
        ((1140, 620, 1450, 710), "D6", "thong_bao"),
        ((1140, 770, 1450, 860), "D10", "activity_logs"),
    ]
    for box, code, label in stores:
        datastore(draw, box, label, code)
    rounded_box(draw, (1490, 855, 1740, 955), "Dịch vụ Email", "#F1F5F9", "#64748B", text_font=FONT_18)
    arrow(draw, (330, 235), (510, 205), "yêu cầu đăng ký", label_offset=(-30, -35))
    arrow(draw, (675, 260), (675, 385), "sự kiện chọn", label_offset=(20, -10))
    arrow(draw, (675, 495), (675, 620), "hợp lệ", label_offset=(20, -10))
    arrow(draw, (675, 730), (675, 855), "kết quả", label_offset=(20, -10))
    arrow(draw, (840, 205), (1140, 195), "đọc sự kiện", label_offset=(20, -35))
    arrow(draw, (840, 440), (1140, 335), "kiểm tra trùng", label_offset=(10, -45))
    arrow(draw, (840, 455), (1140, 475), "thông tin SV", label_offset=(10, 8))
    arrow(draw, (840, 675), (1140, 335), "ghi/cập nhật", label_offset=(0, -120))
    arrow(draw, (840, 675), (1140, 665), "tạo thông báo", label_offset=(0, -35))
    arrow(draw, (840, 675), (1140, 815), "ghi log", label_offset=(0, 40))
    arrow(draw, (840, 910), (1490, 905), "email xác nhận", label_offset=(100, -35))
    img.save(path)


def make_checkin_diagram(path: Path) -> None:
    img, draw = new_canvas("DFD mức 2 - Điểm danh QR và cộng điểm", 1800, 1120)
    rounded_box(draw, (70, 190, 330, 290), "Sinh viên", "#F1F5F9", "#64748B", text_font=FONT_20)
    rounded_box(draw, (70, 770, 330, 870), "Quản trị viên", "#F1F5F9", "#64748B", text_font=FONT_20)
    steps = [
        ((500, 145, 840, 255), "4.1", "Quét QR đầu buổi"),
        ((500, 370, 840, 480), "4.2", "Xác thực đăng ký và thời gian"),
        ((500, 595, 840, 705), "4.3", "Ghi nhận điểm danh"),
        ((500, 820, 840, 930), "4.4", "Tính và cộng điểm rèn luyện"),
    ]
    for box, number, label in steps:
        process_box(draw, box, number, label)
    stores = [
        ((1140, 145, 1460, 235), "D2", "su_kien"),
        ((1140, 285, 1460, 375), "D3", "dang_ky"),
        ((1140, 425, 1460, 515), "D4", "chi_tiet_diem_danh"),
        ((1140, 640, 1460, 730), "D5", "lich_su_diem"),
        ((1140, 790, 1460, 880), "D10", "activity_logs"),
    ]
    for box, code, label in stores:
        datastore(draw, box, label, code)
    arrow(draw, (330, 240), (500, 200), "QR đầu/cuối", label_offset=(-20, -35))
    arrow(draw, (670, 255), (670, 370), "mã QR", label_offset=(20, -8))
    arrow(draw, (670, 480), (670, 595), "hợp lệ", label_offset=(20, -8))
    arrow(draw, (670, 705), (670, 820), "đủ điều kiện", label_offset=(20, -8))
    arrow(draw, (330, 820), (500, 875), "cấu hình điểm", label_offset=(-20, 10))
    arrow(draw, (840, 205), (1140, 190), "đọc sự kiện", label_offset=(15, -35))
    arrow(draw, (840, 425), (1140, 330), "đọc đăng ký", label_offset=(10, -45))
    arrow(draw, (840, 650), (1140, 470), "ghi lượt điểm danh", label_offset=(0, -80))
    arrow(draw, (840, 875), (1140, 685), "ghi lịch sử điểm", label_offset=(0, -85))
    arrow(draw, (840, 875), (1140, 835), "ghi log", label_offset=(10, -35))
    img.save(path)


def generate_figures() -> dict[str, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    figures = {
        "Sơ đồ DFD 4.1": FIG_DIR / "dfd_context_powerdesigner_style.png",
        "Sơ đồ DFD 4.2": FIG_DIR / "dfd_level1_powerdesigner_style.png",
        "Sơ đồ DFD 4.3": FIG_DIR / "dfd_registration_powerdesigner_style.png",
        "Sơ đồ DFD 4.4": FIG_DIR / "dfd_checkin_powerdesigner_style.png",
    }
    make_context_diagram(figures["Sơ đồ DFD 4.1"])
    make_level1_diagram(figures["Sơ đồ DFD 4.2"])
    make_registration_diagram(figures["Sơ đồ DFD 4.3"])
    make_checkin_diagram(figures["Sơ đồ DFD 4.4"])
    return figures


def paragraph_has_drawing_element(element) -> bool:
    return element is not None and "w:drawing" in element.xml


def replace_text_with_highlight(paragraph, text: str, color: WD_COLOR_INDEX) -> None:
    paragraph._element.clear_content()
    run = paragraph.add_run(text)
    run.font.highlight_color = color


def replace_dfd_images(doc: Document, figures: dict[str, Path]) -> int:
    captions = {
        "Sơ đồ DFD 4.1": "Sơ đồ DFD 4.1. Sơ đồ ngữ cảnh hệ thống quản lý sự kiện (dạng PowerDesigner)",
        "Sơ đồ DFD 4.2": "Sơ đồ DFD 4.2. DFD mức 1 cho các tiến trình chính (dạng PowerDesigner)",
        "Sơ đồ DFD 4.3": "Sơ đồ DFD 4.3. DFD mức 2 cho quy trình đăng ký / hủy đăng ký sự kiện (dạng PowerDesigner)",
        "Sơ đồ DFD 4.4": "Sơ đồ DFD 4.4. DFD mức 2 cho quy trình điểm danh QR và cộng điểm (dạng PowerDesigner)",
    }
    replaced = 0
    for key, path in figures.items():
        paragraphs = list(doc.paragraphs)
        caption = next((p for p in paragraphs if key in p.text), None)
        if caption is None:
            continue
        prev = caption._element.getprevious()
        if paragraph_has_drawing_element(prev):
            prev.getparent().remove(prev)
        image_para = caption.insert_paragraph_before("")
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = image_para.add_run()
        run.add_picture(str(path), width=Inches(6.5))
        replace_text_with_highlight(caption, captions[key], WD_COLOR_INDEX.YELLOW)
        replaced += 1
    return replaced


def in_range(index: int, start: int | None, end: int | None) -> bool:
    return start is not None and end is not None and start < index < end


def first_index(paragraphs: list, needles: tuple[str, ...]) -> int | None:
    for i, p in enumerate(paragraphs):
        text = p.text.strip().upper()
        if any(needle.upper() in text for needle in needles):
            return i
    return None


def is_heading_like(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    upper = stripped.upper()
    prefixes = (
        "CHƯƠNG ",
        "CHUONG ",
        "1.",
        "2.",
        "3.",
        "4.",
        "5.",
        "6.",
        "7.",
        "BẢNG ",
        "BANG ",
        "HÌNH ",
        "HINH ",
        "SƠ ĐỒ ",
        "SO DO ",
        "MỤC ",
        "DANH MỤC",
        "TÀI LIỆU",
    )
    return upper.startswith(prefixes) or len(stripped) < 80


def has_numeric_citation(text: str) -> bool:
    for n in range(1, 30):
        if f"[{n}]" in text:
            return True
    return False


def highlight_paragraph(paragraph, color: WD_COLOR_INDEX) -> None:
    if not paragraph.runs and paragraph.text:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        if run.text:
            run.font.highlight_color = color


def append_marker(paragraph, marker: str, color: WD_COLOR_INDEX) -> bool:
    if marker in paragraph.text:
        return False
    run = paragraph.add_run(marker)
    run.font.highlight_color = color
    return True


def apply_dense_text_red(doc: Document) -> int:
    paragraphs = list(doc.paragraphs)
    ch1 = first_index(paragraphs, ("CHƯƠNG 1", "CHUONG 1"))
    ch2 = first_index(paragraphs, ("CHƯƠNG 2", "CHUONG 2"))
    ch3 = first_index(paragraphs, ("CHƯƠNG 3", "CHUONG 3"))
    ch4 = first_index(paragraphs, ("CHƯƠNG 4", "CHUONG 4"))
    count = 0

    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text or is_heading_like(text):
            continue

        mark = False
        if in_range(idx, ch1, ch2) and len(text) >= 120:
            mark = True
        elif in_range(idx, ch2, ch3) and len(text) >= 105:
            mark = True
        elif in_range(idx, ch3, ch4) and len(text) >= 115:
            mark = True

        if mark:
            highlight_paragraph(paragraph, WD_COLOR_INDEX.RED)
            count += 1

    return count


TECH_SOURCE_TERMS = (
    "Laravel",
    "PHP",
    "MySQL",
    "Bootstrap",
    "JavaScript",
    "AJAX",
    "Laravel Sanctum",
    "RESTful API",
    "API",
    "MVC",
    "ORM",
    "HTTP",
    "JSON",
    "SMTP",
    "CSRF",
    "XSS",
    "Docker",
    "React Native",
    "Expo",
    "Axios",
)


def apply_pink_source_markers(doc: Document) -> int:
    paragraphs = list(doc.paragraphs)
    ch1 = first_index(paragraphs, ("CHƯƠNG 1", "CHUONG 1"))
    ch2 = first_index(paragraphs, ("CHƯƠNG 2", "CHUONG 2"))
    ch3 = first_index(paragraphs, ("CHƯƠNG 3", "CHUONG 3"))
    ch4 = first_index(paragraphs, ("CHƯƠNG 4", "CHUONG 4"))
    marker = " [cần trích nguồn]"
    count = 0

    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text or is_heading_like(text) or has_numeric_citation(text):
            continue

        need_marker = False
        if in_range(idx, ch1, ch4) and len(text) >= 105:
            need_marker = True
        elif any(term in text for term in TECH_SOURCE_TERMS) and len(text) >= 90:
            need_marker = True
        elif any(phrase in text for phrase in ("bảo mật", "hiệu năng", "quy trình", "nghiệp vụ", "ứng dụng web")) and len(text) >= 105:
            need_marker = True

        if need_marker and append_marker(paragraph, marker, WD_COLOR_INDEX.PINK):
            count += 1

    return count


def audit_highlights(doc: Document) -> dict[str, int]:
    counts: dict[str, int] = {}
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            color = run.font.highlight_color
            if color is not None:
                key = str(color)
                counts[key] = counts.get(key, 0) + 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        color = run.font.highlight_color
                        if color is not None:
                            key = str(color)
                            counts[key] = counts.get(key, 0) + 1
    return counts


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)

    figures = generate_figures()
    doc = Document(str(SOURCE_DOCX))
    replaced_dfd = replace_dfd_images(doc, figures)
    red_marks = apply_dense_text_red(doc)
    pink_marks = apply_pink_source_markers(doc)

    doc.save(str(WORKSPACE_OUT))
    doc.save(str(DESKTOP_OUT))

    reopened = Document(str(WORKSPACE_OUT))
    counts = audit_highlights(reopened)
    print(f"workspace_out={WORKSPACE_OUT}")
    print(f"desktop_out={DESKTOP_OUT}")
    print(f"replaced_dfd={replaced_dfd}")
    print(f"red_dense_text_marks={red_marks}")
    print(f"pink_source_markers={pink_marks}")
    print(f"paragraphs={len(reopened.paragraphs)} tables={len(reopened.tables)} inline_shapes={len(reopened.inline_shapes)}")
    print("highlight_counts=" + repr(counts))


if __name__ == "__main__":
    main()
