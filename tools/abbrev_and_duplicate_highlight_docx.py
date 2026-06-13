from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SOURCE = Path(r"C:\Users\phuqu\OneDrive\Máy tính\64131942_DuongPhuQuang_DoAn - Sao chép - chinh sua highlight.docx")
WORKSPACE_OUTPUT = Path(r"D:\laragon\www\ql_su_kien\64131942_DuongPhuQuang_DoAn_viet_tat_trung_lap_highlight.docx")
DESKTOP_OUTPUT = Path(r"C:\Users\phuqu\OneDrive\Máy tính\64131942_DuongPhuQuang_DoAn - Sao chép - viet tat trung lap highlight.docx")

BLUE = WD_COLOR_INDEX.TURQUOISE
RED = WD_COLOR_INDEX.RED


ABBREVIATIONS = [
    # Longest and most specific replacements first.
    ("giao diện người dùng và trải nghiệm người dùng", "UI/UX"),
    ("Giao diện người dùng và trải nghiệm người dùng", "UI/UX"),
    ("giao diện và trải nghiệm người dùng", "UI/UX"),
    ("Giao diện và trải nghiệm người dùng", "UI/UX"),
    ("Công nghệ Thông tin", "CNTT"),
    ("Công nghệ thông tin", "CNTT"),
    ("công nghệ thông tin", "CNTT"),
    ("Mã số sinh viên", "MSSV"),
    ("mã số sinh viên", "MSSV"),
    ("cơ sở dữ liệu", "CSDL"),
    ("Cơ sở dữ liệu", "CSDL"),
    ("mã QR", "QR"),
    ("Mã QR", "QR"),
]


DUPLICATE_PREFIXES = [
    # Chapter 5 repeats requirements/business flows already covered in Ch. 3/4.
    "Module người dùng hỗ trợ các chức năng đăng nhập",
    "Module quản lý sự kiện là module trung tâm",
    "Hệ thống hỗ trợ tìm kiếm sự kiện theo tên",
    "Module đăng ký sự kiện cho phép sinh viên đăng ký",
    "Module này giúp kiểm soát chặt chẽ số lượng sinh viên",
    "Module điểm danh QR được xây dựng nhằm thay thế",
    "Quy trình điểm danh có phân biệt điểm danh đầu buổi",
    "Trong trường hợp sinh viên chưa đăng ký nhưng quét QR",
    "Module điểm rèn luyện quản lý việc cộng hoặc trừ điểm",
    "Khi sinh viên tham gia sự kiện và hoàn thành đủ điều kiện",
    "Module thông báo dùng để gửi thông tin đến sinh viên",
    "Sinh viên có thể xem danh sách thông báo",
    "Module thư viện media cho phép quản trị viên upload",
    "Module bầu cử trực tuyến phục vụ các hoạt động bỏ phiếu",
    "Hệ thống quản lý danh sách ứng cử viên",
    "Module báo cáo và thống kê cung cấp số liệu tổng quan",
    "Giao diện sinh viên được xây dựng bằng Blade",
    "Các sự kiện được trình bày với thông tin cơ bản",
    "Giao diện quản trị viên cung cấp các chức năng quản lý",
    "Giao diện quản lý sự kiện cho phép quản trị viên tạo sự kiện",
    "Giao diện điểm danh hỗ trợ quản trị viên theo dõi",
    "Giao diện báo cáo và thống kê hiển thị các số liệu",
    "Ngoài ra, hệ thống hỗ trợ xuất báo cáo ra Excel",
    "Màn hình đăng nhập cho phép sinh viên nhập tài khoản",
    "Màn hình hồ sơ hiển thị thông tin cá nhân",
    "Màn hình danh sách sự kiện cho phép sinh viên xem",
    "Màn hình chi tiết sự kiện hiển thị đầy đủ nội dung",
    "Mobile app cung cấp màn hình thông báo",
    "Màn hình điểm rèn luyện cho phép sinh viên xem tổng điểm",
    # The bridge summary now duplicates the detailed Ch. 6 testing chapter.
    "Sau khi hoàn thành các module chính, hệ thống được kiểm thử",
    # Conclusion repeats implementation and advantages already covered earlier.
    "Về phía backend, hệ thống được xây dựng bằng Laravel",
    "Về phía frontend web, hệ thống cung cấp giao diện cho sinh viên",
    "Bên cạnh giao diện web, đề tài cũng xây dựng ứng dụng mobile",
    "Ngoài ra, hệ thống đã tích hợp một số chức năng mở rộng",
    "Ưu điểm đầu tiên của hệ thống là hỗ trợ tương đối đầy đủ",
    "Thứ hai, hệ thống có sự phân quyền rõ ràng",
    "Thứ ba, việc tích hợp điểm danh bằng QR",
    "Thứ tư, hệ thống có cả web và mobile app",
    "Thứ năm, hệ thống được xây dựng theo kiến trúc khá rõ ràng",
]


def clone_run_format(src, dst):
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    dst.style = src.style
    dst.font.name = src.font.name
    dst.font.size = src.font.size
    dst.font.color.rgb = src.font.color.rgb
    dst.font.highlight_color = src.font.highlight_color
    dst.font.all_caps = src.font.all_caps
    dst.font.small_caps = src.font.small_caps
    dst.font.subscript = src.font.subscript
    dst.font.superscript = src.font.superscript


def clear_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def replace_terms_in_paragraph(paragraph) -> int:
    if not paragraph.runs:
        return 0
    original_runs = [(run, run.text) for run in paragraph.runs]
    changes = 0
    new_segments = []

    for run, text in original_runs:
        segments = [(text, None)]
        for old, new in ABBREVIATIONS:
            next_segments = []
            for seg_text, color in segments:
                if color is not None or old not in seg_text:
                    next_segments.append((seg_text, color))
                    continue
                parts = seg_text.split(old)
                for idx, part in enumerate(parts):
                    if part:
                        next_segments.append((part, None))
                    if idx < len(parts) - 1:
                        next_segments.append((new, BLUE))
                        changes += 1
                segments = next_segments
            segments = next_segments
        new_segments.append((run, segments))

    if changes == 0:
        return 0

    clear_runs(paragraph)
    for old_run, segments in new_segments:
        for text, color in segments:
            if not text:
                continue
            new_run = paragraph.add_run(text)
            clone_run_format(old_run, new_run)
            if color is not None:
                new_run.font.highlight_color = color
    return changes


def should_process_body_paragraph(paragraph, after_intro: bool, before_refs: bool) -> bool:
    if not after_intro or not before_refs:
        return False
    if paragraph.style.name.startswith("Heading"):
        return False
    if paragraph.style.name in {"Caption", "table of figures"}:
        return False
    text = paragraph.text.strip()
    if not text:
        return False
    if text.startswith("[") and "]" in text[:5]:
        return False
    return True


def highlight_duplicate_paragraph(paragraph) -> bool:
    text = " ".join(paragraph.text.split())
    if any(text.startswith(prefix) for prefix in DUPLICATE_PREFIXES):
        for run in paragraph.runs:
            run.font.highlight_color = RED
        return True
    return False


def process_doc(doc: Document) -> tuple[int, int]:
    after_intro = False
    before_refs = True
    abbreviation_changes = 0
    duplicate_marks = 0

    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if text == "LỜI MỞ ĐẦU":
            after_intro = True
        if text == "TÀI LIỆU THAM KHẢO":
            before_refs = False

        if should_process_body_paragraph(paragraph, after_intro, before_refs):
            abbreviation_changes += replace_terms_in_paragraph(paragraph)
        if before_refs and highlight_duplicate_paragraph(paragraph):
            duplicate_marks += 1

    # Process table body cells except the abbreviation table and references-like rows.
    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "Từ viết tắt" in table_text and "Thuật ngữ đầy đủ" in table_text:
            continue
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text.startswith("[") and "]" in text[:5]:
                        continue
                    abbreviation_changes += replace_terms_in_paragraph(paragraph)

    return abbreviation_changes, duplicate_marks


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(str(SOURCE))
    abbreviation_changes, duplicate_marks = process_doc(doc)
    doc.save(str(WORKSPACE_OUTPUT))
    doc.save(str(DESKTOP_OUTPUT))
    print(f"workspace={WORKSPACE_OUTPUT}")
    print(f"desktop={DESKTOP_OUTPUT}")
    print(f"abbreviation_changes={abbreviation_changes}")
    print(f"duplicate_marks={duplicate_marks}")


if __name__ == "__main__":
    main()
