from __future__ import annotations

import math
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path(r"D:\laragon\www\ql_su_kien")
SOURCE = Path(r"C:\Users\phuqu\OneDrive\Máy tính\64131942_DuongPhuQuang_DoAn - Sao chép.docx")
OUTPUT = WORKSPACE / "64131942_DuongPhuQuang_DoAn_chinh_sua_highlight.docx"
FIG_DIR = WORKSPACE / "docx_generated_figures"
YELLOW = WD_COLOR_INDEX.YELLOW


def get_font(size: int, bold: bool = False):
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap_text(text: str, font, max_width: int, draw: ImageDraw.ImageDraw) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else current + " " + word
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, text, font, fill="#ffffff", outline="#2563eb", text_fill="#111827", width=3):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=width)
    lines = wrap_text(text, font, x2 - x1 - 24, draw)
    line_h = font.size + 6
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=font, fill=text_fill)
        y += line_h


def draw_arrow(draw, start, end, color="#334155", width=3):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    angle = math.atan2(y2 - y1, x2 - x1)
    head = 15
    for delta in (math.pi / 7, -math.pi / 7):
        x = x2 - head * math.cos(angle + delta)
        y = y2 - head * math.sin(angle + delta)
        draw.line((x2, y2, x, y), fill=color, width=width)


def draw_label(draw, text, xy, font, fill="#334155"):
    x, y = xy
    lines = text.split("\n")
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + 4


def save_context_dfd(path: Path):
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = get_font(38, True)
    font = get_font(28, True)
    small = get_font(22)
    draw.text((50, 35), "DFD mức ngữ cảnh - Hệ thống quản lý sự kiện", font=title_font, fill="#0f172a")
    center = (560, 290, 1040, 560)
    draw_box(draw, center, "Hệ thống quản lý sự kiện", get_font(34, True), fill="#eff6ff", outline="#1d4ed8")
    actors = {
        "Sinh viên": (80, 170, 360, 300),
        "Quản trị viên": (80, 610, 360, 740),
        "Mobile app": (1200, 170, 1480, 300),
        "SMTP/Email": (1200, 610, 1480, 740),
    }
    for text, box in actors.items():
        draw_box(draw, box, text, font, fill="#ffffff", outline="#64748b")
    draw_arrow(draw, (360, 235), (560, 365))
    draw_arrow(draw, (560, 430), (360, 675))
    draw_arrow(draw, (1200, 235), (1040, 365))
    draw_arrow(draw, (1040, 430), (1200, 675))
    draw_label(draw, "Đăng ký, quét QR,\nxem điểm, bỏ phiếu", (385, 235), small)
    draw_label(draw, "Quản lý sự kiện,\nbáo cáo, cấu hình", (390, 610), small)
    draw_label(draw, "REST API,\nJSON, token", (1070, 235), small)
    draw_label(draw, "Email xác thực,\nđặt lại mật khẩu", (1060, 610), small)
    img.save(path)


def save_level1_dfd(path: Path):
    img = Image.new("RGB", (1800, 1050), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = get_font(36, True)
    font = get_font(22, True)
    small = get_font(18)
    draw.text((45, 30), "DFD mức 1 - Các tiến trình chính và kho dữ liệu", font=title_font, fill="#0f172a")
    processes = [
        ("1. Xác thực", 430, 120),
        ("2. Quản lý sự kiện", 430, 230),
        ("3. Đăng ký/Hủy", 430, 340),
        ("4. Điểm danh QR", 430, 450),
        ("5. Điểm rèn luyện", 430, 560),
        ("6. Thông báo/Email", 430, 670),
        ("7. Báo cáo/Thống kê", 430, 780),
        ("8. Bầu cử/Media/Chatbot", 430, 890),
    ]
    stores = [
        ("D1 Người dùng", 1120, 120),
        ("D2 Sự kiện", 1120, 230),
        ("D3 Đăng ký", 1120, 340),
        ("D4 Điểm danh", 1120, 450),
        ("D5 Điểm", 1120, 560),
        ("D6 Thông báo", 1120, 670),
        ("D7 Bầu cử", 1120, 780),
        ("D8 Media, D9 Cấu hình", 1120, 890),
    ]
    draw_box(draw, (70, 260, 290, 390), "Sinh viên", font, fill="#ffffff", outline="#64748b")
    draw_box(draw, (70, 610, 290, 740), "Quản trị viên", font, fill="#ffffff", outline="#64748b")
    for text, x, y in processes:
        draw_box(draw, (x, y, x + 430, y + 78), text, font, fill="#eff6ff", outline="#1d4ed8")
    for text, x, y in stores:
        draw_box(draw, (x, y, x + 360, y + 78), text, font, fill="#fff7ed", outline="#ea580c")
    for _, x, y in processes:
        draw_arrow(draw, (290, 325), (x, y + 39), width=2)
        draw_arrow(draw, (290, 675), (x, y + 39), width=2)
    for i, (_, x, y) in enumerate(processes):
        sx, sy = stores[i][1], stores[i][2]
        draw_arrow(draw, (x + 430, y + 39), (sx, sy + 39), color="#475569", width=2)
    draw_label(draw, "Yêu cầu nghiệp vụ", (120, 410), small)
    draw_label(draw, "Quản trị và cấu hình", (95, 755), small)
    img.save(path)


def save_registration_dfd(path: Path):
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = get_font(36, True)
    font = get_font(22, True)
    draw.text((45, 30), "DFD mức 2 - Đăng ký và hủy đăng ký sự kiện", font=title_font, fill="#0f172a")
    steps = [
        "3.1 Nhận yêu cầu",
        "3.2 Kiểm tra đăng nhập",
        "3.3 Kiểm tra sự kiện",
        "3.4 Kiểm tra sức chứa/trùng",
        "3.5 Ghi đăng ký",
        "3.6 Trả kết quả",
    ]
    x = 170
    y = 230
    prev = None
    for i, step in enumerate(steps):
        box = (x + i * 230, y, x + i * 230 + 190, y + 105)
        draw_box(draw, box, step, font, fill="#eff6ff", outline="#1d4ed8")
        if prev:
            draw_arrow(draw, (prev[2], y + 52), (box[0], y + 52), width=2)
        prev = box
    stores = [
        ("D1 Người dùng", 300, 550),
        ("D2 Sự kiện", 620, 550),
        ("D3 Đăng ký", 940, 550),
        ("D6 Thông báo", 1220, 550),
    ]
    for text, sx, sy in stores:
        draw_box(draw, (sx, sy, sx + 230, sy + 90), text, font, fill="#fff7ed", outline="#ea580c")
    draw_arrow(draw, (465, 335), (410, 550), width=2)
    draw_arrow(draw, (775, 335), (735, 550), width=2)
    draw_arrow(draw, (1005, 335), (1060, 550), width=2)
    draw_arrow(draw, (1360, 335), (1330, 550), width=2)
    img.save(path)


def save_checkin_dfd(path: Path):
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = get_font(36, True)
    font = get_font(20, True)
    draw.text((45, 30), "DFD mức 2 - Điểm danh QR và cộng điểm", font=title_font, fill="#0f172a")
    steps = [
        "4.1 Nhận QR",
        "4.2 Kiểm tra token",
        "4.3 Xác định sự kiện/đăng ký",
        "4.4 Ghi điểm danh",
        "4.5 Kiểm tra đủ điều kiện",
        "4.6 Cộng điểm/thông báo",
    ]
    coords = [(55, 190), (300, 190), (545, 190), (790, 190), (1035, 190), (1280, 190)]
    boxes = []
    for step, (x, y) in zip(steps, coords):
        box = (x, y, x + 215, y + 105)
        boxes.append(box)
        draw_box(draw, box, step, font, fill="#eff6ff", outline="#1d4ed8")
    for a, b in zip(boxes, boxes[1:]):
        draw_arrow(draw, (a[2], (a[1] + a[3]) // 2), (b[0], (b[1] + b[3]) // 2), width=2)
    stores = [
        ("D2 Sự kiện", 230, 560),
        ("D3 Đăng ký", 485, 560),
        ("D4 Điểm danh", 740, 560),
        ("D5 Điểm", 995, 560),
        ("D6 Thông báo", 1250, 560),
    ]
    for text, x, y in stores:
        draw_box(draw, (x, y, x + 220, y + 88), text, font, fill="#fff7ed", outline="#ea580c")
    # Vertical data-store links stay below the process lane to avoid crossing labels.
    draw_arrow(draw, (652, 295), (340, 560), width=2)
    draw_arrow(draw, (760, 295), (595, 560), width=2)
    draw_arrow(draw, (897, 295), (850, 560), width=2)
    draw_arrow(draw, (1388, 295), (1105, 560), width=2)
    draw_arrow(draw, (1480, 295), (1360, 560), width=2)
    img.save(path)


def generate_figures() -> dict[str, Path]:
    FIG_DIR.mkdir(exist_ok=True)
    paths = {
        "context": FIG_DIR / "dfd_context.png",
        "level1": FIG_DIR / "dfd_level1.png",
        "registration": FIG_DIR / "dfd_registration.png",
        "checkin": FIG_DIR / "dfd_checkin.png",
    }
    save_context_dfd(paths["context"])
    save_level1_dfd(paths["level1"])
    save_registration_dfd(paths["registration"])
    save_checkin_dfd(paths["checkin"])
    return paths


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph, text: str, highlight: bool = False):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    if highlight:
        run.font.highlight_color = YELLOW
    return run


def highlight_paragraph(paragraph):
    for run in paragraph.runs:
        run.font.highlight_color = YELLOW


def insert_highlight_para_before(target, text: str, style: str | None = None, bold: bool = False):
    para = target.insert_paragraph_before("", style=style)
    run = para.add_run(text)
    run.font.highlight_color = YELLOW
    run.bold = bold
    return para


def append_highlight_run(paragraph, text: str):
    if text in paragraph.text:
        return
    run = paragraph.add_run(text)
    run.font.highlight_color = YELLOW
    return run


def find_paragraph(doc: Document, needle: str, style_prefix: str | None = None):
    for p in doc.paragraphs:
        if needle.lower() in p.text.lower():
            if style_prefix is None or p.style.name.startswith(style_prefix):
                return p
    raise ValueError(f"Cannot find paragraph containing {needle!r}")


def find_exact_heading(doc: Document, text: str, style_prefix: str):
    normalized = " ".join(text.split()).lower()
    for p in doc.paragraphs:
        if p.style.name.startswith(style_prefix) and " ".join(p.text.split()).lower() == normalized:
            return p
    raise ValueError(f"Cannot find heading {text!r}")


def add_table_before(doc: Document, target, rows: list[list[str]], widths: list[float] | None = None, font_size: float = 9.0):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.highlight_color = YELLOW
            run.font.size = Pt(font_size)
            if r_idx == 0:
                run.bold = True
            if widths:
                cell.width = Inches(widths[c_idx])
    target._p.addprevious(table._tbl)
    return table


def add_image_before(target, image_path: Path, width_inches: float = 6.4):
    para = target.insert_paragraph_before("")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return para


def remove_between(doc: Document, start_para, end_para):
    paragraphs = doc.paragraphs
    elements = [p._element for p in paragraphs]
    start = elements.index(start_para._element)
    end = elements.index(end_para._element)
    for p in paragraphs[start + 1 : end]:
        p._element.getparent().remove(p._element)


def append_citation(doc: Document, startswith: str, citation: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            append_highlight_run(p, f" {citation}")
            return
    raise ValueError(f"Cannot find paragraph starting with {startswith!r}")


def replace_all_text_fragment(doc: Document, old: str, new: str):
    for p in doc.paragraphs:
        if old in p.text:
            set_paragraph_text(p, p.text.replace(old, new), highlight=True)


def insert_abbreviations(doc: Document):
    target = find_paragraph(doc, "LỜI MỞ ĐẦU")
    intro = insert_highlight_para_before(
        target,
        "Bảng dưới đây bổ sung các ký hiệu và từ viết tắt được sử dụng nhiều trong báo cáo.",
    )
    rows = [
        ["Từ viết tắt", "Thuật ngữ đầy đủ", "Giải thích trong báo cáo"],
        ["API", "Application Programming Interface", "Giao diện lập trình ứng dụng, dùng cho mobile app và các request dữ liệu."],
        ["QR", "Quick Response", "Mã QR dùng cho điểm danh sự kiện và QR cá nhân sinh viên."],
        ["MVC", "Model - View - Controller", "Mô hình tổ chức ứng dụng Laravel."],
        ["ORM", "Object Relational Mapping", "Cơ chế ánh xạ model với bảng dữ liệu, trong dự án dùng Eloquent ORM."],
        ["REST", "Representational State Transfer", "Phong cách thiết kế API cho web/mobile."],
        ["HTTP", "HyperText Transfer Protocol", "Giao thức request/response giữa client và server."],
        ["JSON", "JavaScript Object Notation", "Định dạng dữ liệu API trả về cho mobile app."],
        ["SMTP", "Simple Mail Transfer Protocol", "Giao thức gửi email xác thực, đặt lại mật khẩu và thông báo."],
        ["CSRF", "Cross-Site Request Forgery", "Kiểu tấn công giả mạo request; Laravel có middleware chống CSRF."],
        ["XSS", "Cross-Site Scripting", "Kiểu tấn công chèn mã độc vào nội dung hiển thị."],
        ["UI/UX", "User Interface/User Experience", "Giao diện người dùng và trải nghiệm người dùng."],
        ["CSDL", "Cơ sở dữ liệu", "Nơi lưu trữ dữ liệu nghiệp vụ của hệ thống."],
        ["CNTT", "Công nghệ thông tin", "Khoa Công nghệ Thông tin."],
        ["MSSV", "Mã số sinh viên", "Mã định danh sinh viên, đồng thời là khóa chính của bảng nguoi_dung."],
    ]
    add_table_before(doc, target, rows, widths=[1.0, 2.2, 3.1], font_size=8.5)
    insert_highlight_para_before(target, "")


def fix_intro_and_citations(doc: Document):
    # Correct technology versions against composer.json and the project overview.
    p = find_paragraph(doc, "Laravel 11")
    new_text = (
        "Hệ thống được phát triển trên nền Laravel 10, PHP 8.1 trở lên; môi trường Docker của dự án sử dụng PHP 8.2. "
        "Cơ sở dữ liệu dùng MySQL 8.0, giao diện web dùng Blade kết hợp Tailwind CSS, tài nguyên frontend được build bằng Vite, "
        "API REST được bảo vệ bởi Laravel Sanctum và hệ thống được đóng gói triển khai bằng Docker Compose [4], [5], [10], [11], [12], [15]. "
        "Các chức năng trọng tâm gồm: quản lý sự kiện và phân loại; đăng ký/hủy; điểm danh bằng QR và ghi nhận thời gian; tính điểm rèn luyện; "
        "gửi thông báo; xuất báo cáo; và cung cấp ứng dụng mobile cho sinh viên."
    )
    set_paragraph_text(p, new_text, highlight=True)
    append_citation(doc, "Đề tài kế thừa và áp dụng quy trình", "[2], [14]")
    append_citation(doc, "Kiến trúc hệ thống web là", "[2]")
    append_citation(doc, "Mô hình MVC gồm ba phần", "[2]")
    append_citation(doc, "Trong hệ thống quản lý sự kiện, lớp Route", "[4]")
    append_citation(doc, "Trong hệ thống quản trị sự kiện, cơ sở dữ liệu quan hệ", "[1], [3], [12]")
    append_citation(doc, "Thiết kế cơ sở dữ liệu cần tuân", "[1], [3]")
    append_citation(doc, "Khóa chính là thuộc tính", "[1], [3]")
    append_citation(doc, "PHP là ngôn ngữ lập trình", "[7]")
    append_citation(doc, "Laravel là một framework PHP", "[4]")
    append_citation(doc, "MySQL là hệ quản trị cơ sở dữ liệu", "[12]")
    append_citation(doc, "Eloquent ORM là công cụ", "[4]")
    append_citation(doc, "Vite là công cụ hỗ trợ build", "[10]")
    append_citation(doc, "Tailwind CSS là framework CSS", "[11]")
    append_citation(doc, "Laravel Sanctum là thư viện", "[5]")
    append_citation(doc, "Composer là công cụ quản lý", "[8]")
    append_citation(doc, "Node.js là môi trường chạy", "[9]")
    append_citation(doc, "Docker Compose là công cụ", "[15]")
    append_citation(doc, "PHPUnit là framework kiểm thử", "[14]")
    append_citation(doc, "Trước khi triển khai, tài nguyên frontend", "[10]")
    append_citation(doc, "Dự án có cấu hình Docker Compose", "[15]")
    append_citation(doc, "Service app chạy Laravel/PHP-FPM", "[12], [13]")
    append_citation(doc, "Thứ tư, chức năng thông báo trong hệ thống", "[6]")


def renumber_chapter4(doc: Document):
    replacements = [
        ("4.3 THIẾT KẾ CƠ SỞ DỮ LIỆU", "4.4 THIẾT KẾ CƠ SỞ DỮ LIỆU"),
        ("4.4 THIẾT KẾ KIẾN TRÚC HỆ THỐNG", "4.5 THIẾT KẾ KIẾN TRÚC HỆ THỐNG"),
        ("4.5 THIẾT KẾ GIAO DIỆN", "4.6 THIẾT KẾ GIAO DIỆN"),
        ("4.6 THIẾT KẾ BẢO MẬT VÀ PHÂN QUYỀN", "4.7 THIẾT KẾ BẢO MẬT VÀ PHÂN QUYỀN"),
    ]
    for old, new in replacements:
        try:
            p = find_paragraph(doc, old, "Heading")
            set_paragraph_text(p, new, highlight=True)
        except ValueError:
            pass

    for p in doc.paragraphs:
        text = p.text.strip()
        if not p.style.name.startswith("Heading 3"):
            continue
        if text.startswith("4.3."):
            set_paragraph_text(p, "4.4." + text[len("4.3.") :], highlight=True)
        elif text.startswith("4.4."):
            set_paragraph_text(p, "4.5." + text[len("4.4.") :], highlight=True)
        elif text.startswith("4.5."):
            set_paragraph_text(p, "4.6." + text[len("4.5.") :], highlight=True)
        elif text.startswith("4.6."):
            set_paragraph_text(p, "4.7." + text[len("4.6.") :], highlight=True)


def insert_dfd_section(doc: Document, figures: dict[str, Path]):
    target = find_paragraph(doc, "4.4 THIẾT KẾ CƠ SỞ DỮ LIỆU", "Heading")
    insert_highlight_para_before(target, "4.3 THIẾT KẾ LUỒNG DỮ LIỆU DFD", "Heading 2")
    insert_highlight_para_before(
        target,
        "Sơ đồ luồng dữ liệu (DFD) được bổ sung để làm rõ cách dữ liệu di chuyển giữa sinh viên, quản trị viên, các tiến trình xử lý và các kho dữ liệu chính. "
        "Phần này giúp liên kết trực tiếp các yêu cầu nghiệp vụ ở Chương 3 với thiết kế cơ sở dữ liệu ở mục 4.4.",
    )
    insert_highlight_para_before(target, "4.3.1 DFD mức ngữ cảnh", "Heading 3")
    add_image_before(target, figures["context"], 6.5)
    insert_highlight_para_before(target, "Sơ đồ DFD 4.1. DFD mức ngữ cảnh của hệ thống quản lý sự kiện", "Caption")
    insert_highlight_para_before(
        target,
        "Ở mức ngữ cảnh, hệ thống quản lý sự kiện là trung tâm tiếp nhận yêu cầu từ sinh viên, quản trị viên và mobile app. "
        "Hệ thống trả về dữ liệu sự kiện, thông báo, điểm rèn luyện, kết quả bầu cử và báo cáo; đồng thời kết nối SMTP để gửi email xác thực, đặt lại mật khẩu hoặc thông báo.",
    )
    insert_highlight_para_before(target, "4.3.2 DFD mức 1 của hệ thống", "Heading 3")
    add_image_before(target, figures["level1"], 6.5)
    insert_highlight_para_before(target, "Sơ đồ DFD 4.2. DFD mức 1 của hệ thống", "Caption")
    rows = [
        ["Kho dữ liệu", "Bảng tương ứng", "Vai trò"],
        ["D1 Người dùng", "nguoi_dung, personal_access_tokens, password_reset_tokens", "Lưu tài khoản, vai trò, token và dữ liệu xác thực."],
        ["D2 Sự kiện", "su_kien, loai_su_kien, mau_bai_dang", "Lưu thông tin sự kiện, loại sự kiện và bố cục bài đăng."],
        ["D3 Đăng ký", "dang_ky", "Lưu lượt đăng ký/hủy đăng ký và trạng thái tham gia."],
        ["D4 Điểm danh", "chi_tiet_diem_danh", "Lưu các lần điểm danh đầu buổi/cuối buổi."],
        ["D5 Điểm", "lich_su_diem", "Lưu lịch sử cộng/trừ điểm rèn luyện."],
        ["D6 Thông báo", "thong_bao, lich_gui_thong_bao, smtp_settings", "Lưu thông báo và thông tin gửi email/thông báo."],
        ["D7 Bầu cử", "bau_cu, ung_cu_vien, cu_tri, phieu_bau, chi_tiet_phieu_bau", "Lưu dữ liệu bầu cử và kết quả bỏ phiếu."],
        ["D8 Media, D9 Cấu hình", "thu_vien_da_phuong_tien, activity_logs, gemini_settings", "Lưu file, nhật ký hoạt động và cấu hình mở rộng."],
    ]
    add_table_before(doc, target, rows, widths=[1.25, 2.8, 2.35], font_size=8.0)
    insert_highlight_para_before(target, "4.3.3 DFD mức 2 cho đăng ký và hủy đăng ký sự kiện", "Heading 3")
    add_image_before(target, figures["registration"], 6.5)
    insert_highlight_para_before(target, "Sơ đồ DFD 4.3. DFD mức 2 cho đăng ký và hủy đăng ký sự kiện", "Caption")
    insert_highlight_para_before(
        target,
        "Luồng đăng ký sử dụng dữ liệu người dùng để kiểm tra đăng nhập, dữ liệu sự kiện để kiểm tra thời gian và sức chứa, dữ liệu đăng ký để chống đăng ký trùng, sau đó ghi nhận kết quả vào bảng dang_ky và tạo thông báo phản hồi cho sinh viên.",
    )
    insert_highlight_para_before(target, "4.3.4 DFD mức 2 cho điểm danh QR", "Heading 3")
    add_image_before(target, figures["checkin"], 6.5)
    insert_highlight_para_before(target, "Sơ đồ DFD 4.4. DFD mức 2 cho điểm danh QR và cộng điểm", "Caption")
    insert_highlight_para_before(
        target,
        "Luồng điểm danh QR giải mã dữ liệu QR, kiểm tra sự kiện và lượt đăng ký, ghi chi tiết điểm danh vào chi_tiet_diem_danh, sau đó kiểm tra điều kiện cộng điểm. Khi sinh viên đủ số lần điểm danh yêu cầu, hệ thống ghi lịch sử điểm vào lich_su_diem và tạo thông báo cập nhật điểm.",
    )


def replace_55_summary(doc: Document):
    test_heading = find_exact_heading(doc, "5.5 Kiểm thử hệ thống", "Heading 2")
    deploy_heading = find_exact_heading(doc, "5.6 Triển khai hệ thống", "Heading 2")
    remove_between(doc, test_heading, deploy_heading)
    set_paragraph_text(test_heading, "5.5 Tổng quan kiểm thử", highlight=True)
    insert_highlight_para_before(
        deploy_heading,
        "Sau khi hoàn thành các module chính, hệ thống được kiểm thử trên cả giao diện web, REST API và ứng dụng mobile. "
        "Nội dung kiểm thử tập trung vào các nghiệp vụ cốt lõi như xác thực, phân quyền, quản lý sự kiện, đăng ký/hủy đăng ký, điểm danh QR, cộng điểm rèn luyện, thông báo, bầu cử và báo cáo. "
        "Quy trình kiểm thử, dữ liệu kiểm thử, test case và kết quả chi tiết được trình bày trong Chương 6.",
    )


def insert_chapter6(doc: Document):
    conclusion = find_exact_heading(doc, "KẾT LUẬN", "Heading 1")
    insert_highlight_para_before(conclusion, "KIỂM THỬ HỆ THỐNG WEB VÀ MOBILE APP", "Heading 1")
    insert_highlight_para_before(
        conclusion,
        "Chương này trình bày quy trình kiểm thử hệ thống sau khi hoàn thành phần xây dựng ứng dụng. Nội dung kiểm thử được tổ chức theo hướng kiểm tra chức năng, API, nghiệp vụ tích hợp, bảo mật cơ bản, giao diện web/mobile và kiểm thử hồi quy.",
    )

    insert_highlight_para_before(conclusion, "6.1 Mục tiêu kiểm thử", "Heading 2")
    insert_highlight_para_before(
        conclusion,
        "Mục tiêu kiểm thử là xác nhận hệ thống đáp ứng các yêu cầu đã xác định ở Chương 3 và phù hợp với thiết kế ở Chương 4. Trọng tâm kiểm thử gồm xác thực, phân quyền, quản lý sự kiện, đăng ký/hủy đăng ký, điểm danh QR, cộng điểm rèn luyện, thông báo, media, bầu cử, báo cáo, API mobile và khả năng hiển thị trên web/mobile [2], [14].",
    )

    insert_highlight_para_before(conclusion, "6.2 Phạm vi kiểm thử", "Heading 2")
    rows = [
        ["Phạm vi", "Nội dung kiểm thử"],
        ["Web sinh viên", "Đăng nhập, xem/tìm kiếm sự kiện, đăng ký, hủy đăng ký, lịch sử tham gia, thông báo, bầu cử."],
        ["Web quản trị", "Quản lý người dùng, sự kiện, điểm danh, media, bầu cử, thống kê, báo cáo, SMTP, activity logs."],
        ["REST API", "Auth API, event API, registration API, QR API, notification API, point API, voting API."],
        ["Mobile app", "Đăng nhập, danh sách/chi tiết sự kiện, QR scanner, thông báo, điểm rèn luyện, bầu cử, hồ sơ cá nhân."],
        ["Ngoài phạm vi kiểm thử sâu", "Kiểm thử tải lớn production, kiểm thử xâm nhập chuyên sâu và kiểm thử toàn bộ thiết bị mobile thực tế."],
    ]
    add_table_before(doc, conclusion, rows, widths=[1.7, 4.7], font_size=8.5)

    insert_highlight_para_before(conclusion, "6.3 Môi trường và công cụ kiểm thử", "Heading 2")
    rows = [
        ["Thành phần", "Môi trường/công cụ"],
        ["Backend", "Laravel 10, PHP ^8.1; Docker runtime dùng PHP 8.2."],
        ["Cơ sở dữ liệu", "MySQL 8 hoặc database test tương đương."],
        ["Web frontend", "Blade, Vite 5, trình duyệt Chrome/Edge."],
        ["Mobile", "React Native 0.83.6, Expo 55, thiết bị hoặc emulator Android."],
        ["API", "Laravel Sanctum, HTTP client/Postman hoặc Laravel Feature Test."],
        ["Tự động hóa", "PHPUnit 10, Laravel Unit Test và Feature Test."],
        ["Kiểm tra build", "npm run build, php artisan test."],
    ]
    add_table_before(doc, conclusion, rows, widths=[1.7, 4.7], font_size=8.5)

    insert_highlight_para_before(conclusion, "6.4 Dữ liệu kiểm thử", "Heading 2")
    rows = [
        ["Nhóm dữ liệu", "Dữ liệu cần chuẩn bị"],
        ["Người dùng", "1 admin hoạt động, 2 sinh viên đã xác thực email, 1 sinh viên chưa xác thực, 1 tài khoản bị khóa."],
        ["Sự kiện", "Sự kiện còn chỗ, sự kiện đã đầy, sự kiện đã bắt đầu, sự kiện đã kết thúc, sự kiện bị hủy."],
        ["Đăng ký", "Đăng ký hợp lệ, đăng ký trùng, đăng ký đã hủy, đăng ký đã điểm danh một lần, đăng ký đủ hai lần điểm danh."],
        ["QR", "QR đầu buổi, QR cuối buổi, QR hết hạn, QR sai định dạng, QR cá nhân sinh viên."],
        ["Bầu cử", "Cuộc bầu cử đang diễn ra, đã kết thúc, có cử tri hợp lệ, sinh viên không thuộc danh sách cử tri."],
        ["Media/Báo cáo", "Ảnh hợp lệ, file sai định dạng, lớp có dữ liệu điểm, lớp không có dữ liệu, khoảng thời gian hợp lệ/không hợp lệ."],
    ]
    add_table_before(doc, conclusion, rows, widths=[1.7, 4.7], font_size=8.5)

    insert_highlight_para_before(conclusion, "6.5 Chiến lược kiểm thử", "Heading 2")
    rows = [
        ["Loại kiểm thử", "Mục tiêu"],
        ["Kiểm thử đơn vị", "Kiểm tra service xử lý nghiệp vụ như đăng ký, hủy đăng ký, điểm danh và thống kê sự kiện."],
        ["Kiểm thử chức năng", "Kiểm tra từng chức năng trên web và mobile theo yêu cầu."],
        ["Kiểm thử API", "Kiểm tra status code, JSON response, validation, token và phân quyền."],
        ["Kiểm thử tích hợp", "Kiểm tra luồng đăng ký -> điểm danh -> cộng điểm -> thông báo."],
        ["Kiểm thử bảo mật cơ bản", "Kiểm tra session, token, vai trò, CSRF, validation và giới hạn upload."],
        ["Kiểm thử giao diện", "Kiểm tra tiếng Việt, responsive, vùng an toàn mobile, lỗi ảnh/storage."],
        ["Kiểm thử hồi quy", "Chạy lại các test sau khi sửa lỗi để bảo đảm chức năng cũ không bị ảnh hưởng."],
    ]
    add_table_before(doc, conclusion, rows, widths=[1.7, 4.7], font_size=8.5)

    insert_highlight_para_before(conclusion, "6.6 Ma trận test case", "Heading 2")
    test_cases = [
        ["TC-01", "Xác thực", "Admin đăng nhập đúng thông tin", "Vào được trang quản trị."],
        ["TC-02", "Xác thực", "Sinh viên đăng nhập đúng thông tin", "Vào được trang chủ sinh viên."],
        ["TC-03", "Xác thực", "Đăng nhập sai mật khẩu", "Báo lỗi, không tạo session."],
        ["TC-04", "API Auth", "Sinh viên đã xác thực đăng nhập mobile", "Trả token Sanctum và thông tin user."],
        ["TC-05", "API Auth", "Sinh viên chưa xác thực email đăng nhập API", "Từ chối đăng nhập hoặc yêu cầu xác thực."],
        ["TC-06", "Phân quyền", "Sinh viên truy cập route admin", "Bị chuyển hướng hoặc từ chối."],
        ["TC-07", "Phân quyền API", "Sinh viên gọi API admin", "API trả 403."],
        ["TC-08", "Người dùng", "Admin tạo sinh viên mới", "Tạo user, gửi email xác thực khi SMTP sẵn sàng."],
        ["TC-09", "Người dùng", "Admin lọc người dùng theo lớp", "Danh sách chỉ hiển thị sinh viên thuộc lớp đã chọn."],
        ["TC-10", "Sự kiện", "Admin tạo sự kiện hợp lệ", "Sự kiện được lưu và hiển thị."],
        ["TC-11", "Sự kiện", "Tạo sự kiện trùng lịch/địa điểm", "Cảnh báo hoặc từ chối theo rule."],
        ["TC-12", "Sự kiện", "Admin cập nhật sự kiện", "Dữ liệu được cập nhật đúng."],
        ["TC-13", "Sự kiện", "Admin xóa sự kiện", "Sự kiện bị xóa mềm hoặc không còn hiển thị."],
        ["TC-14", "API Sự kiện", "Mobile lấy danh sách sự kiện", "Trả JSON đúng cấu trúc."],
        ["TC-15", "API Sự kiện", "Tìm kiếm sự kiện", "Trả sự kiện phù hợp từ khóa."],
        ["TC-16", "Đăng ký", "Sinh viên đăng ký sự kiện còn chỗ", "Tạo dang_ky, tăng số lượng hiện tại."],
        ["TC-17", "Đăng ký", "Đăng ký trùng", "Từ chối đăng ký trùng."],
        ["TC-18", "Đăng ký", "Đăng ký sự kiện đã đầy", "Từ chối và báo hết chỗ."],
        ["TC-19", "Đăng ký", "Đăng ký sự kiện đã bắt đầu", "Từ chối đăng ký."],
        ["TC-20", "Hủy đăng ký", "Hủy trước giờ diễn ra", "Hủy đăng ký, giảm số lượng hiện tại."],
        ["TC-21", "Hủy đăng ký", "Hủy sau khi sự kiện bắt đầu", "Từ chối hủy."],
        ["TC-22", "QR", "Quét QR đầu buổi hợp lệ", "Ghi chi_tiet_diem_danh loại dau_buoi."],
        ["TC-23", "QR", "Quét QR cuối buổi hợp lệ", "Ghi cuoi_buoi, cập nhật đủ điều kiện."],
        ["TC-24", "QR", "Quét trùng cùng loại", "Báo đã điểm danh, không ghi trùng."],
        ["TC-25", "QR", "Quét QR sai hoặc hết hạn", "Từ chối, không ghi điểm danh."],
        ["TC-26", "QR", "Chưa đăng ký nhưng quét QR", "Tạo đăng ký hoặc từ chối theo rule thiết kế."],
        ["TC-27", "Điểm", "Đủ điều kiện cộng điểm", "Tạo lich_su_diem đúng số điểm, không cộng lặp."],
        ["TC-28", "Điểm", "Admin cộng/trừ điểm thủ công", "Ghi lịch sử điểm đúng nguồn."],
        ["TC-29", "Thông báo", "Sinh viên xem thông báo", "Hiển thị danh sách và badge chưa đọc."],
        ["TC-30", "Thông báo", "Đánh dấu đã đọc", "Cập nhật da_doc."],
        ["TC-31", "Media", "Upload ảnh hợp lệ", "File lưu đúng storage, hiển thị trong thư viện."],
        ["TC-32", "Media", "Upload file sai định dạng", "Từ chối bằng validation."],
        ["TC-33", "Bầu cử", "Cử tri hợp lệ bỏ phiếu", "Tạo phiếu bầu, đánh dấu đã bỏ phiếu."],
        ["TC-34", "Bầu cử", "Sinh viên không phải cử tri bỏ phiếu", "Từ chối bỏ phiếu."],
        ["TC-35", "Báo cáo", "Xuất báo cáo lớp có dữ liệu", "Xuất Excel đúng dữ liệu."],
        ["TC-36", "Báo cáo", "Xuất báo cáo không có dữ liệu", "Báo không có dữ liệu, không xuất file rỗng."],
        ["TC-37", "Giao diện", "Kiểm tra responsive web", "Không vỡ layout, chữ không chồng lấn."],
        ["TC-38", "Mobile", "Kiểm tra vùng an toàn và font tiếng Việt", "Không mất nội dung, tiếng Việt hiển thị đúng."],
        ["TC-39", "Bảo mật", "POST form web không có CSRF token", "Request bị từ chối."],
        ["TC-40", "Hồi quy", "Chạy lại test tự động sau sửa lỗi", "Các test quan trọng pass."],
    ]
    rows = [["ID", "Nhóm", "Kịch bản", "Kết quả mong đợi"]] + test_cases
    add_table_before(doc, conclusion, rows, widths=[0.7, 1.0, 2.6, 2.2], font_size=7.6)

    insert_highlight_para_before(conclusion, "6.7 Kết quả kiểm thử theo nhóm chức năng", "Heading 2")
    rows = [
        ["Nhóm", "Test tự động/kiểm thử liên quan", "Kết quả cần trình bày"],
        ["API xác thực", "tests/Feature/Api/AuthApiTest.php", "Đăng nhập mobile, tài khoản chưa xác thực, đăng ký mobile, gửi lại email xác thực."],
        ["API sự kiện", "tests/Feature/Api/EventApiTest.php", "Danh sách, chi tiết, tìm kiếm, tạo/sửa/xóa sự kiện."],
        ["API đăng ký", "tests/Feature/Api/RegistrationApiTest.php", "Đăng ký, chống đăng ký trùng, hủy, lịch sử, admin scan QR."],
        ["QR", "tests/Feature/Api/QrCodeApiTest.php; tests/Feature/CheckinScannerTest.php", "Tạo QR SVG/PNG, scanner, QR hết hạn, QR cá nhân sinh viên."],
        ["Service nghiệp vụ", "tests/Unit/Services/EventServiceTest.php; RegistrationServiceTest.php", "Tạo/sửa sự kiện, trùng lịch, đăng ký, hủy, điểm danh, cộng điểm một lần."],
        ["Admin/Chatbot", "tests/Feature/Admin/*; tests/Feature/Chatbot/GeminiChatbotTest.php", "Activity logs, báo cáo, cấu hình Gemini, quản lý người dùng, chatbot fallback/cache."],
    ]
    add_table_before(doc, conclusion, rows, widths=[1.4, 2.4, 2.6], font_size=7.8)

    insert_highlight_para_before(conclusion, "6.8 Lỗi phát hiện và hướng xử lý", "Heading 2")
    rows = [
        ["Mã lỗi", "Nhóm", "Mô tả", "Hướng xử lý", "Trạng thái"],
        ["BUG-01", "Storage", "Ảnh upload không hiển thị do đường dẫn storage/link public.", "Kiểm tra storage:link và chuẩn hóa đường dẫn lưu file.", "Đã khắc phục"],
        ["BUG-02", "Mobile API", "Mobile app không kết nối đúng backend do sai base URL/IP.", "Cấu hình EXPO_PUBLIC_API_URL đúng IP máy chạy Laravel.", "Đã khắc phục"],
        ["BUG-03", "Giao diện", "Font tiếng Việt hiển thị chưa ổn định ở một số màn hình.", "Chuẩn hóa font và kiểm tra render trên web/mobile.", "Đã khắc phục"],
        ["BUG-04", "Mobile UI", "Một số nội dung bị sát vùng status bar/notch.", "Dùng SafeAreaView/safe area context.", "Đã khắc phục"],
        ["BUG-05", "Upload avatar", "Upload avatar lỗi validation hoặc đường dẫn.", "Rà rule upload, storage path và hiển thị ảnh.", "Đã khắc phục"],
    ]
    add_table_before(doc, conclusion, rows, widths=[0.75, 1.1, 2.2, 2.1, 1.0], font_size=7.4)

    insert_highlight_para_before(conclusion, "6.9 Đánh giá sau kiểm thử", "Heading 2")
    insert_highlight_para_before(
        conclusion,
        "Qua kiểm thử, hệ thống đáp ứng được các nghiệp vụ cốt lõi của một hệ thống quản lý sự kiện trong môi trường đại học. "
        "Các luồng quan trọng như đăng nhập, phân quyền, quản lý sự kiện, đăng ký/hủy đăng ký, điểm danh QR, cộng điểm rèn luyện, thông báo, bầu cử và báo cáo có thể hoạt động trên web và mobile. "
        "Hệ thống vẫn cần tiếp tục mở rộng kiểm thử tự động, kiểm thử hiệu năng, kiểm thử bảo mật chuyên sâu và kiểm thử trên nhiều thiết bị mobile thực tế trước khi triển khai ở quy mô lớn.",
    )


def renumber_conclusion(doc: Document):
    conclusion = find_exact_heading(doc, "KẾT LUẬN", "Heading 1")
    references = find_exact_heading(doc, "TÀI LIỆU THAM KHẢO", "Heading 1")
    set_paragraph_text(conclusion, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", highlight=True)
    paragraphs = doc.paragraphs
    elements = [p._element for p in paragraphs]
    start = elements.index(conclusion._element)
    end = elements.index(references._element)
    for p in paragraphs[start + 1 : end]:
        if not p.style.name.startswith("Heading 2"):
            continue
        text = p.text.strip()
        if text.startswith("6."):
            set_paragraph_text(p, "7." + text[len("6.") :], highlight=True)


def enable_update_fields(docx_path: Path):
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                from lxml import etree

                root = etree.fromstring(data)
                existing = root.find(qn("w:updateFields"))
                if existing is None:
                    update = OxmlElement("w:updateFields")
                    update.set(qn("w:val"), "true")
                    root.append(update)
                else:
                    existing.set(qn("w:val"), "true")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            zout.writestr(item, data)
    tmp.replace(docx_path)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    figures = generate_figures()
    doc = Document(str(SOURCE))

    replace_all_text_fragment(doc, "Khánh Hoà", "Khánh Hòa")
    insert_abbreviations(doc)
    fix_intro_and_citations(doc)
    renumber_chapter4(doc)
    insert_dfd_section(doc, figures)
    replace_55_summary(doc)
    insert_chapter6(doc)
    renumber_conclusion(doc)

    doc.save(str(OUTPUT))
    enable_update_fields(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
